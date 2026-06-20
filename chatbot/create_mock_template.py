import os
from docx import Document

def create_mock():
    os.makedirs("templates", exist_ok=True)
    doc = Document()
    doc.add_heading('THƯ CẢNH BÁO VI PHẠM', 0)
    
    doc.add_paragraph('Kính gửi: {{TEN_DOI_TAC}}')
    doc.add_paragraph('Địa chỉ: {{DIA_CHI}}')
    
    doc.add_paragraph('Chúng tôi là công ty {{TEN_CONG_TY}}. Qua quá trình theo dõi, chúng tôi phát hiện quý công ty đã có hành vi vi phạm: {{HANH_VI_VI_PHAM}}.')
    
    doc.add_paragraph('Hành vi này vi phạm trực tiếp đến quyền lợi hợp pháp của chúng tôi. Yêu cầu quý công ty chấm dứt ngay hành vi trên trước ngày {{NGAY_HET_HAN}}.')
    
    doc.add_paragraph('Trân trọng,')
    doc.add_paragraph('Đại diện pháp luật')
    
    doc.save("templates/Mau_Thu_Canh_Bao.docx")
    print("Created templates/Mau_Thu_Canh_Bao.docx")

if __name__ == "__main__":
    create_mock()
