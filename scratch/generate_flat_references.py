import docx
from docx.shared import Inches, Pt, RGBColor
import os
import shutil

doc = docx.Document()

# Adjust margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Add title
title = doc.add_paragraph()
title_run = title.add_run("TÀI LIỆU THAM KHẢO")
title_run.bold = True
title_run.font.size = Pt(14)
title_run.font.name = "Arial"
title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

# Add spacing
p_space = doc.add_paragraph()
p_space.paragraph_format.space_after = Pt(12)

refs = [
    ("Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., ... & Kiela, D. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. ", "Advances in Neural Information Processing Systems (NeurIPS)", ", 33, 9459-9474."),
    ("Gao, Y., Xiong, Y., Huang, S., et al. (2024). Retrieval-Augmented Generation for Large Language Models: A Survey. ", "arXiv preprint arXiv:2312.10997", "."),
    ("Cormack, G. V., Clarke, C. L., & Buettcher, S. (2009). Reciprocal Rank Fusion Outperforms Condorcet and Individual Rank Learning Methods. ", "Proceedings of the 32nd International ACM SIGIR Conference", ", 758-759."),
    ("Robertson, S. E., & Zaragoza, H. (2009). The Probabilistic Relevance Framework: BM25 and Beyond. ", "Foundations and Trends in Information Retrieval", ", 3(4), 333-389."),
    ("Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. ", "Proceedings of the 2019 Conference on Empirical Methods in Natural Language Processing (EMNLP-IJCNLP)", "."),
    ("Guo, Z., Jin, R., Liu, C., et al. (2024). LightRAG: Simple and Fast Retrieval-Augmented Generation. ", "arXiv preprint arXiv:2410.05779", "."),
    ("Nogueira, R., & Cho, K. (2019). Passage Re-ranking with BERT. ", "arXiv preprint arXiv:1901.04085", "."),
    ("Qdrant. (2024). ", "Qdrant Vector Search Engine Documentation", ". https://qdrant.tech/documentation/"),
    ("FastAPI. (2024). ", "FastAPI Framework Documentation", ". https://fastapi.tiangolo.com/"),
    ("Quốc Hội Nước CHXHCN Việt Nam. (2020). ", "Luật Doanh nghiệp số 59/2020/QH14", "."),
    ("Quốc Hội Nước CHXHCN Việt Nam. (2020). ", "Luật Đầu tư số 61/2020/QH14", "."),
    ("Quốc Hội Nước CHXHCN Việt Nam. (2019). ", "Bộ Luật Lao động số 45/2019/QH14", "."),
    ("Quốc Hội Nước CHXHCN Việt Nam. (2005). ", "Luật Thương mại số 36/2005/QH11", "."),
    ("Quốc Hội Nước CHXHCN Việt Nam. (2019). ", "Luật Quản lý thuế số 38/2019/QH14", "."),
    ("Chính Phủ Nước CHXHCN Việt Nam. (2020). ", "Nghị định số 123/2020/NĐ-CP quy định về hóa đơn, chứng từ", "."),
    ("Chính Phủ Nước CHXHCN Việt Nam. (2021). ", "Nghị định số 31/2021/NĐ-CP quy định chi tiết và hướng dẫn thi hành một số điều của Luật Đầu tư", "."),
    ("Bộ Tài chính. (2016). ", "Thông tư số 133/2016/TT-BTC hướng dẫn Chế độ kế toán doanh nghiệp nhỏ và vừa", ".")
]

for idx, (prefix, italic_part, suffix) in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after = Pt(6)
    
    # [idx]
    r_num = p.add_run(f"[{idx}] ")
    r_num.font.name = "Arial"
    r_num.font.size = Pt(10)
    
    # Prefix
    r_pre = p.add_run(prefix)
    r_pre.font.name = "Arial"
    r_pre.font.size = Pt(10)
    
    # Italic part
    r_ital = p.add_run(italic_part)
    r_ital.font.name = "Arial"
    r_ital.font.size = Pt(10)
    r_ital.italic = True
    
    # Suffix
    r_suf = p.add_run(suffix)
    r_suf.font.name = "Arial"
    r_suf.font.size = Pt(10)

doc.save("references.docx")
shutil.copy("references.docx", "references.doc")
print("Done")
