import os
import requests
from bs4 import BeautifulSoup
from docx import Document
import logging
import re
import pathlib

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("TemplateCrawler")

class TemplateCrawler:
    def __init__(self):
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        self.base_dir = pathlib.Path(__file__).parent.parent
        self.template_dir = self.base_dir / "templates"
        self.template_dir.mkdir(exist_ok=True)

    def fetch_html(self, url):
        try:
            logger.info(f"Đang crawl dữ liệu từ: {url}")
            response = requests.get(url, headers=self.headers, timeout=15)
            response.raise_for_status()
            response.encoding = 'utf-8'
            return response.text
        except Exception as e:
            logger.error(f"Lỗi khi tải trang {url}: {e}")
            return None

    def extract_contract_text(self, html_content, source_type="luatvietnam"):
        """Trích xuất phần văn bản của hợp đồng từ HTML HTML"""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Tùy thuộc vào source, cấu trúc HTML sẽ khác nhau
        if source_type == "luatvietnam":
            content_div = soup.find('div', class_='article-content') or soup.find('div', id='doc-content')
        elif source_type == "thuvienphapluat":
            content_div = soup.find('div', id='divContentDoc') or soup.find('div', class_='content1')
        elif source_type == "hoatieu":
            content_div = soup.find('div', id='noidung') or soup.find('div', class_='detail-content')
        else:
            content_div = soup.find('article') or soup.body
            
        if not content_div:
            # Fallback lấy toàn bộ text nếu không tìm thấy div chính xác
            logger.warning("Không tìm thấy thẻ div chứa nội dung chuẩn, chuyển sang chế độ fallback.")
            paragraphs = soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'div', 'td'])
        else:
            paragraphs = content_div.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'div', 'td'])
            
        text_lines = []
        for p in paragraphs:
            text = p.get_text(strip=True)
            if text and len(text) > 2:
                # Bỏ qua các text rác như quảng cáo, link tải
                if any(x in text.lower() for x in ['tải về', 'download', 'bản quyền', 'quảng cáo']):
                    continue
                text_lines.append(text)
                
        # Lọc bỏ trùng lặp và nối lại
        seen = set()
        unique_lines = []
        for line in text_lines:
            if line not in seen:
                seen.add(line)
                unique_lines.append(line)
                
        return "\n".join(unique_lines)

    def process_placeholders(self, text):
        """Tự động nhận diện các đoạn cần điền và chuyển thành format {{placeholder}}"""
        # Thay thế "........." hoặc "___" thành các trường dữ liệu
        text = re.sub(r'\.{5,}', ' {{thong_tin}} ', text)
        text = re.sub(r'_{5,}', ' {{thong_tin}} ', text)
        
        # Nhận diện tên công ty, cá nhân
        text = re.sub(r'Ông/Bà:\s*{{thong_tin}}', 'Ông/Bà: {{ten_ca_nhan}}', text, flags=re.IGNORECASE)
        text = re.sub(r'Công ty:\s*{{thong_tin}}', 'Công ty: {{ten_cong_ty}}', text, flags=re.IGNORECASE)
        text = re.sub(r'Địa chỉ:\s*{{thong_tin}}', 'Địa chỉ: {{dia_chi}}', text, flags=re.IGNORECASE)
        text = re.sub(r'Mã số thuế:\s*{{thong_tin}}', 'Mã số thuế: {{ma_so_thue}}', text, flags=re.IGNORECASE)
        
        return text

    def save_to_docx(self, text, filename):
        """Lưu văn bản vào file docx với định dạng chuẩn"""
        doc = Document()
        
        # Set font chuẩn Times New Roman
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        
        lines = text.split('\n')
        for line in lines:
            if "CỘNG HÒA XÃ HỘI" in line.upper() or "ĐỘC LẬP" in line.upper():
                p = doc.add_paragraph(line)
                p.alignment = 1 # Center
                p.runs[0].bold = True
            elif line.isupper() and len(line) < 100:
                # Tiêu đề hợp đồng
                p = doc.add_paragraph(line)
                p.alignment = 1 # Center
                p.runs[0].bold = True
            elif line.startswith("Điều") or line.startswith("ĐIỀU"):
                p = doc.add_paragraph(line)
                p.runs[0].bold = True
            else:
                doc.add_paragraph(line)
                
        output_path = self.template_dir / f"{filename}.docx"
        doc.save(str(output_path))
        logger.info(f"Đã lưu template tại: {output_path}")
        return output_path

    def crawl_and_save(self, url, filename, source_type="luatvietnam"):
        html = self.fetch_html(url)
        if not html:
            return False
            
        raw_text = self.extract_contract_text(html, source_type)
        if len(raw_text) < 500:
            logger.warning(f"Văn bản trích xuất quá ngắn ({len(raw_text)} ký tự). Có thể crawler không lấy được đúng nội dung.")
            
        processed_text = self.process_placeholders(raw_text)
        self.save_to_docx(processed_text, filename)
        return True

if __name__ == "__main__":
    crawler = TemplateCrawler()
    
    # Danh sách các URL cần crawl (có thể mở rộng thêm theo yêu cầu của SME)
    targets = [
        # Nhóm Dân sự & Thương mại
        {"url": "https://luatvietnam.vn/bieu-mau/hop-dong-mua-ban-hang-hoa-182-27156-article.html", "filename": "Hop_dong_mua_ban_hang_hoa", "source": "luatvietnam"},
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-dich-vu-moi-nhat-182-23456-article.html", "filename": "Hop_dong_dich_vu", "source": "luatvietnam"},
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-thue-nha-182-27158-article.html", "filename": "Hop_dong_thue_nha", "source": "luatvietnam"},
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-vay-tien-182-27159-article.html", "filename": "Hop_dong_vay_tien", "source": "luatvietnam"},
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-dai-ly-182-27160-article.html", "filename": "Hop_dong_dai_ly", "source": "luatvietnam"},
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-uy-quyen-182-27161-article.html", "filename": "Hop_dong_uy_quyen", "source": "luatvietnam"},
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-gia-cong-182-27162-article.html", "filename": "Hop_dong_gia_cong", "source": "luatvietnam"},
        
        # Nhóm Doanh nghiệp & Đầu tư
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-hop-tac-kinh-doanh-bcc-182-27163-article.html", "filename": "Hop_dong_hop_tac_kinh_doanh_BCC", "source": "luatvietnam"},
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-lien-ket-182-27164-article.html", "filename": "Hop_dong_lien_ket", "source": "luatvietnam"},
        
        # Nhóm Lao động
        {"url": "https://luatvietnam.vn/bieu-mau/hop-dong-lao-dong-thoi-vu-moi-nhat-182-23456-article.html", "filename": "Hop_dong_lao_dong", "source": "luatvietnam"},
        
        # Nhóm Tài chính - Ngân hàng
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-tin-dung-182-27165-article.html", "filename": "Hop_dong_tin_dung", "source": "luatvietnam"},
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-bao-lanh-182-27166-article.html", "filename": "Hop_dong_bao_lanh", "source": "luatvietnam"},
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-the-chap-182-27167-article.html", "filename": "Hop_dong_the_chap", "source": "luatvietnam"},
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-cam-co-182-27168-article.html", "filename": "Hop_dong_cam_co", "source": "luatvietnam"},
        
        # Nhóm Công nghệ & Sở hữu trí tuệ
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-chuyen-giao-cong-nghe-182-27169-article.html", "filename": "Hop_dong_chuyen_giao_cong_nghe", "source": "luatvietnam"},
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-nhuong-quyen-thuong-mai-182-27170-article.html", "filename": "Hop_dong_nhuong_quyen_thuong_mai", "source": "luatvietnam"},
        {"url": "https://luatvietnam.vn/bieu-mau/mau-hop-dong-so-huu-tri-tue-182-27171-article.html", "filename": "Hop_dong_so_huu_tri_tue", "source": "luatvietnam"},
    ]
    
    for target in targets:
        crawler.crawl_and_save(target["url"], target["filename"], target["source"])
