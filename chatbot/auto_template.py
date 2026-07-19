import os
import pathlib
from docx import Document

def fill_template(template_name: str, context: dict) -> str:
    """
    Điền dữ liệu vào file template .docx
    Trả về đường dẫn file output đã được điền.
    """
    base_dir = pathlib.Path(__file__).parent
    template_path = base_dir / "templates" / f"{template_name}.docx"
    
    if not template_path.exists():
        raise FileNotFoundError(f"Không tìm thấy template: {template_name}.docx")
        
    doc = Document(str(template_path))
    
    # Chuẩn hóa context keys về chữ thường và cắt khoảng trắng
    normalized_context = {}
    for key, value in context.items():
        normalized_context[str(key).lower().strip()] = value
        
    # Hàm xử lý thay thế text trong các đoạn văn (hỗ trợ tránh phân mảnh run)
    def replace_in_paragraph(p):
        for key, value in normalized_context.items():
            # Thử các trường hợp viết hoa viết thường khác nhau của placeholder
            placeholders = [
                f"{{{{{key}}}}}",
                f"{{{{{key.upper()}}}}}",
                f"{{{{{key.lower()}}}}}",
                f"{{{{{key.title()}}}}}"
            ]
            for placeholder in placeholders:
                if placeholder in p.text:
                    replaced = False
                    # Thử thay thế trong từng run đơn lẻ
                    for run in p.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
                            replaced = True
                    # Fallback nếu placeholder bị phân mảnh qua nhiều runs khác nhau
                    if not replaced and p.runs:
                        full_text = p.text.replace(placeholder, str(value))
                        p.runs[0].text = full_text
                        for run in p.runs[1:]:
                            run.text = ""

    # Thay thế trong các đoạn văn chính
    for p in doc.paragraphs:
        replace_in_paragraph(p)
                        
    # Thay thế trong bảng
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    output_dir = base_dir / "outputs"
    output_dir.mkdir(exist_ok=True)
    
    import uuid
    output_filename = f"filled_{template_name}_{uuid.uuid4().hex[:8]}.docx"
    output_path = output_dir / output_filename
    
    doc.save(str(output_path))
    return str(output_path)
