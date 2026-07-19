import pathlib
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_templates():
    templates_dir = pathlib.Path(r"d:\Project Save\chatbot law\chatbot\templates")
    templates_dir.mkdir(exist_ok=True)
    log_path = pathlib.Path(r"d:\Project Save\chatbot law\scratch\generate_templates_log.txt")

    # Định nghĩa cấu trúc điều khoản cho 17 hợp đồng
    contracts_data = {
        "Hop_dong_lao_dong": {
            "title": "HỢP ĐỒNG LAO ĐỘNG",
            "clauses": [
                ("Điều 1: Thời hạn và công việc hợp đồng", 
                 "- Vị trí chuyên môn / Chức danh công việc: {{vi_tri_cong_viec}}.\n- Thời hạn hợp đồng: {{thoi_han_hop_dong}} từ ngày {{ngay_bat_dau}}.\n- Địa điểm làm việc: {{dia_diem_lam_viec}}."),
                ("Điều 2: Chế độ làm việc", 
                 "- Thời giờ làm việc: {{thoi_gio_lam_viec}} giờ/tuần.\n- Thời giờ nghỉ ngơi và chế độ nghỉ lễ Tết theo quy định của pháp luật lao động hiện hành."),
                ("Điều 3: Quyền lợi của người lao động", 
                 "- Mức lương chính: {{muc_luong}} VNĐ/tháng.\n- Phụ cấp và các khoản bổ sung khác: {{phu_cap}} VNĐ/tháng.\n- Đóng bảo hiểm xã hội (BHXH), bảo hiểm y tế (BHYT), bảo hiểm thất nghiệp (BHTN) theo quy định."),
                ("Điều 4: Nghĩa vụ của người lao động", 
                 "- Hoàn thành tốt các nhiệm vụ công việc được giao.\n- Chấp hành nghiêm chỉnh nội quy lao động, an toàn lao động và cam kết bảo mật thông tin."),
                ("Điều 5: Quyền hạn và nghĩa vụ của người sử dụng lao động", 
                 "- Đảm bảo việc làm, điều kiện lao động an toàn và thanh toán tiền lương đầy đủ đúng hạn.\n- Có quyền quản lý điều hành và xử lý vi phạm kỷ luật theo quy chế của doanh nghiệp.")
            ]
        },
        "Hop_dong_dich_vu": {
            "title": "HỢP ĐỒNG DỊCH VỤ",
            "clauses": [
                ("Điều 1: Nội dung dịch vụ yêu cầu", 
                 "- Bên B cam kết thực hiện công việc dịch vụ: {{noi_dung_dich_vu}} cho Bên A theo đúng thời hạn, tiến độ và tiêu chuẩn chất lượng đã cam kết."),
                ("Điều 2: Phí dịch vụ và phương thức thanh toán", 
                 "- Tổng giá trị hợp đồng dịch vụ là: {{phi_dich_vu}} VNĐ.\n- Phương thức thanh toán: {{phuong_thuc_thanh_toan}}.\n- Thời hạn thanh toán: {{han_thanh_toan}}."),
                ("Điều 3: Quyền và nghĩa vụ của Bên A", 
                 "- Cung cấp đầy đủ thông tin, tài liệu và tạo điều kiện thuận lợi để Bên B thực hiện công việc.\n- Tổ chức nghiệm thu kết quả và thực hiện nghĩa vụ thanh toán phí dịch vụ đầy đủ, đúng hạn."),
                ("Điều 4: Quyền và nghĩa vụ của Bên B", 
                 "- Thực hiện dịch vụ cẩn trọng, bảo đảm chất lượng và bàn giao kết quả đúng thời hạn.\n- Bảo mật tuyệt đối mọi thông tin liên quan đến hoạt động kinh doanh của Bên A."),
                ("Điều 5: Phạt vi phạm và giải quyết tranh chấp", 
                 "- Mọi tranh chấp phát sinh được ưu tiên giải quyết qua thương lượng, hòa giải.\n- Trường hợp vi phạm tiến độ hoặc chất lượng, bên vi phạm phải chịu phạt vi phạm {{ty_le_phat}}% giá trị hợp đồng.")
            ]
        },
        "Hop_dong_thue_nha": {
            "title": "HỢP ĐỒNG THUÊ NHÀ",
            "clauses": [
                ("Điều 1: Đối tượng thuê và mục đích sử dụng", 
                 "- Bên A đồng ý cho Bên B thuê căn nhà tại địa chỉ: {{dia_chi_nha_thue}}.\n- Diện tích sử dụng: {{dien_tich}} m2.\n- Mục đích sử dụng: {{muc_dich_thue}} (để ở hoặc làm văn phòng)."),
                ("Điều 2: Thời hạn thuê và thời điểm giao nhận", 
                 "- Thời hạn thuê nhà là: {{thoi_han_thue}} từ ngày {{ngay_bat_dau_thue}}.\n- Thời điểm bàn giao nhà thực tế: ngày {{ngay_ban_giao}}."),
                ("Điều 3: Giá thuê và phương thức thanh toán", 
                 "- Giá thuê nhà là: {{gia_thue_nha}} VNĐ/tháng.\n- Tiền đặt cọc bảo đảm thực hiện hợp đồng: {{tien_dat_coc}} VNĐ.\n- Phương thức thanh toán tiền thuê: {{phuong_thuc_thanh_toan}} vào ngày {{ngay_nop_tien}} hàng tháng."),
                ("Điều 4: Quyền và nghĩa vụ của Bên cho thuê (Bên A)", 
                 "- Bàn giao nhà và trang thiết bị kèm theo đúng tình trạng cam kết.\n- Bảo đảm quyền sử dụng nhà độc lập và ổn định cho Bên B trong thời hạn thuê."),
                ("Điều 5: Quyền và nghĩa vụ của Bên thuê (Bên B)", 
                 "- Sử dụng nhà đúng mục đích, giữ gìn tài sản và tự chịu chi phí điện, nước, internet phát sinh.\n- Thanh toán đầy đủ tiền thuê đúng hạn. Không được tự ý cho thuê lại khi chưa có sự đồng ý của Bên A.")
            ]
        },
        "Hop_dong_mua_ban_hang_hoa": {
            "title": "HỢP ĐỒNG MUA BÁN HÀNG HÓA",
            "clauses": [
                ("Điều 1: Tên hàng, số lượng và chất lượng sản phẩm", 
                 "- Tên hàng hóa giao dịch: {{ten_hang_hoa}}.\n- Số lượng hàng hóa: {{so_luong}}.\n- Yêu cầu chất lượng và thông số kỹ thuật: {{chat_luong_hang}}."),
                ("Điều 2: Đơn giá và phương thức thanh toán", 
                 "- Đơn giá hàng hóa: {{don_gia}} VNĐ/sản phẩm.\n- Tổng giá trị đơn hàng: {{gia_tri_don_hang}} VNĐ.\n- Phương thức thanh toán tiền hàng: {{phuong_thuc_thanh_toan}}."),
                ("Điều 3: Thời gian, địa điểm và phương thức giao nhận", 
                 "- Thời hạn giao nhận hàng: trước ngày {{ngay_giao_hang}}.\n- Địa điểm bàn giao hàng: {{dia_diem_giao}}.\n- Chi phí vận chuyển và xếp dỡ hàng hóa do {{bên_chiu_phi}} chịu trách nhiệm."),
                ("Điều 4: Trách nhiệm bảo hành và xử lý lỗi hàng hóa", 
                 "- Thời gian bảo hành sản phẩm: {{thoi_gian_bao_hanh}}.\n- Trường hợp hàng hóa bị lỗi, kém chất lượng hoặc thiếu hụt, Bên B có trách nhiệm đổi trả hoặc bổ sung trong vòng {{han_doi_tra}} ngày."),
                ("Điều 5: Phạt vi phạm hợp đồng và bồi thường thiệt hại", 
                 "- Bên nào chậm trễ thực hiện nghĩa vụ phải chịu phạt chậm trả / chậm giao hàng {{ty_le_phat}}% / ngày chậm trễ.\n- Mọi thiệt hại thực tế phát sinh phải được bồi thường đầy đủ theo luật định.")
            ]
        },
        "Hop_dong_vay_tien": {
            "title": "HỢP ĐỒNG VAY TIỀN",
            "clauses": [
                ("Điều 1: Số tiền vay và thời hạn vay", 
                 "- Bên A đồng ý cho Bên B vay số tiền là: {{so_tien_vay}} VNĐ.\n- Thời hạn vay tiền là: {{thoi_han_vay}} tháng kể từ ngày nhận đủ tiền vay."),
                ("Điều 2: Lãi suất vay thỏa thuận", 
                 "- Mức lãi suất cho vay là: {{lai_suat_vay}} %/năm.\n- Lãi suất quá hạn (nếu có) được tính bằng {{lai_suat_qua_han}} % mức lãi suất trong hạn."),
                ("Điều 3: Phương thức trả nợ gốc và lãi", 
                 "- Tiền lãi được thanh toán hàng tháng vào ngày {{ngay_tra_lai}}.\n- Tiền gốc được trả một lần vào ngày đáo hạn hợp đồng hoặc trả dần theo kỳ thỏa thuận."),
                ("Điều 4: Cam kết bảo đảm tiền vay", 
                 "- Bên B cam kết sử dụng tiền vay đúng mục đích và tự chịu trách nhiệm hoàn trả cả gốc lẫn lãi đúng hạn.\n- Phương thức bảo đảm tiền vay (nếu có): {{tai_san_bao_dam}}."),
                ("Điều 5: Xử lý vi phạm và tranh chấp", 
                 "- Nếu Bên B không thanh toán nợ đúng hạn, Bên A có quyền yêu cầu cơ quan pháp luật xử lý tài sản bảo đảm để thu hồi nợ.")
            ]
        },
        "Hop_dong_dai_ly": {
            "title": "HỢP ĐỒNG ĐẠI LÝ",
            "clauses": [
                ("Điều 1: Chỉ định đại lý và phạm vi đại lý", 
                 "- Bên A chỉ định Bên B làm đại lý phân phối sản phẩm: {{ten_san_pham}} tại khu vực {{khu_vuc_dai_ly}}.\n- Hình thức đại lý: {{hinh_thuc_dai_ly}} (Đại lý độc quyền hoặc đại lý thông thường)."),
                ("Điều 2: Giá cung cấp và hoa hồng đại lý", 
                 "- Giá bán sản phẩm áp dụng cho đại lý: {{gia_cung_cap}} VNĐ.\n- Tỷ lệ hoa hồng chiết khấu Bên B được hưởng: {{ty_le_hoa_hong}} % trên doanh số bán hàng."),
                ("Điều 3: Giao nhận hàng và thanh toán", 
                 "- Bên A giao hàng cho Bên B theo từng đơn đặt hàng cụ thể.\n- Bên B thanh toán tiền hàng định kỳ vào ngày {{ngay_doi_soat}} hàng tháng sau khi đối chiếu công nợ."),
                ("Điều 4: Quyền và nghĩa vụ của Bên A", 
                 "- Cung cấp hàng hóa đúng chất lượng, hỗ trợ tài liệu quảng cáo và đào tạo nghiệp vụ.\n- Kiểm tra định kỳ việc thực hiện chính sách đại lý của Bên B."),
                ("Điều 5: Quyền và nghĩa vụ của Bên B", 
                 "- Bán hàng đúng giá niêm yết, chăm sóc khách hàng tốt tại khu vực được giao.\n- Không kinh doanh các sản phẩm cạnh tranh trực tiếp với sản phẩm của Bên A.")
            ]
        },
        "Hop_dong_uy_quyen": {
            "title": "HỢP ĐỒNG ỦY QUYỀN",
            "clauses": [
                ("Điều 1: Phạm vi công việc ủy quyền", 
                 "- Bên A ủy quyền cho Bên B thực hiện các công việc sau nhân danh Bên A: {{noi_dung_uy_quyen}}."),
                ("Điều 2: Thời hạn ủy quyền", 
                 "- Thời hạn ủy quyền có hiệu lực kể từ ngày ký đến ngày {{ngay_het_han_uy_quyen}} hoặc khi công việc ủy quyền được hoàn thành."),
                ("Điều 3: Phí thù lao và chi phí thực hiện công việc", 
                 "- Thù lao ủy quyền cho Bên B là: {{thu_lao_uy_quyen}} VNĐ.\n- Các chi phí hợp lý phát sinh để thực hiện công việc ủy quyền do Bên A chi trả và cung cấp ứng trước."),
                ("Điều 4: Quyền và nghĩa vụ của Bên ủy quyền (Bên A)", 
                 "- Cung cấp đầy đủ thông tin, hồ sơ, giấy ủy quyền cần thiết.\n- Chịu trách nhiệm về mọi cam kết do Bên B thực hiện trong phạm vi ủy quyền."),
                ("Điều 5: Quyền và nghĩa vụ của Bên nhận ủy quyền (Bên B)", 
                 "- Thực hiện công việc ủy quyền đúng pháp luật và thường xuyên báo cáo tiến độ cho Bên A.\n- Bàn giao đầy đủ giấy tờ, kết quả công việc sau khi hoàn tất.")
            ]
        },
        "Hop_dong_gia_cong": {
            "title": "HỢP ĐỒNG GIA CÔNG",
            "clauses": [
                ("Điều 1: Tên sản phẩm gia công và quy cách kỹ thuật", 
                 "- Bên B nhận gia công sản phẩm: {{ten_san_pham_gia_cong}}.\n- Quy cách, kích thước, tiêu chuẩn chất lượng sản phẩm chi tiết: {{quy_cach_ky_thuat}}."),
                ("Điều 2: Định mức và cấp phát nguyên vật liệu", 
                 "- Nguyên vật liệu chính do Bên {{ben_cung_cap_nguyen_lieu}} cung cấp.\n- Định mức hao hụt cho phép trong quá trình sản xuất gia công: {{dinh_muc_hao_hut}} %."),
                ("Điều 3: Đơn giá gia công và thanh toán", 
                 "- Đơn giá gia công sản phẩm: {{don_gia_gia_cong}} VNĐ/đơn vị.\n- Tổng phí gia công dự kiến: {{tong_phi_gia_cong}} VNĐ.\n- Phương thức thanh toán: {{phuong_thuc_thanh_toan}}."),
                ("Điều 4: Thời gian và địa điểm giao nhận sản phẩm", 
                 "- Thời hạn giao nhận thành phẩm gia công: trước ngày {{ngay_giao_hang}}.\n- Địa điểm giao hàng và nghiệm thu sản phẩm: {{dia_diem_giao}}."),
                ("Điều 5: Trách nhiệm do vi phạm chất lượng", 
                 "- Bên B phải sửa chữa hoặc làm lại sản phẩm nếu không đạt yêu cầu chất lượng.\n- Chịu trách nhiệm bồi thường nguyên vật liệu bị hỏng hoặc hao hụt vượt định mức.")
            ]
        },
        "Hop_dong_hop_tac_kinh_doanh_BCC": {
            "title": "HỢP ĐỒNG HỢP TÁC KINH DOANH (BCC)",
            "clauses": [
                ("Điều 1: Mục tiêu và phạm vi hợp tác", 
                 "- Hai bên đồng ý hợp tác thực hiện dự án kinh doanh: {{ten_du_an}} mà không thành lập pháp nhân mới theo quy định của Luật Đầu tư."),
                ("Điều 2: Phần vốn góp của mỗi bên", 
                 "- Bên A đóng góp vốn bằng: {{gop_von_ben_a}} VNĐ.\n- Bên B đóng góp vốn bằng: {{gop_von_ben_b}} VNĐ (hoặc bằng tài sản/công nghệ/đất đai)."),
                ("Điều 3: Phân chia kết quả kinh doanh và lợi nhuận", 
                 "- Tỷ lệ chia sẻ doanh thu / phân chia lợi nhuận sau thuế: Bên A hưởng {{ty_le_ben_a}} %, Bên B hưởng {{ty_le_ben_b}} %.\n- Thời điểm đối soát và chia lợi nhuận định kỳ: {{ky_chia_loi_nhuan}}."),
                ("Điều 4: Ban điều hành và quản lý hoạt động hợp tác", 
                 "- Thành lập Ban điều hành dự án gồm {{so_luong_nhan_su}} người để quyết định các vấn đề lớn.\n- Người đại diện theo pháp luật điều hành hoạt động hàng ngày: {{nguoi_dai_dien}}."),
                ("Điều 5: Cam kết chung và thời hạn hợp tác", 
                 "- Thời hạn hợp tác kinh doanh theo hợp đồng: {{thoi_han_hop_tac}} năm.\n- Mọi sửa đổi, bổ sung điều khoản hợp đồng phải được hai bên ký văn bản đồng ý.")
            ]
        },
        "Hop_dong_lien_ket": {
            "title": "HỢP ĐỒNG LIÊN KẾT KINH DOANH",
            "clauses": [
                ("Điều 1: Mục đích và nội dung liên kết", 
                 "- Hai bên cùng liên kết triển khai chương trình / hoạt động: {{noi_dung_lien_ket}} nhằm khai thác thế mạnh nguồn lực của mỗi bên."),
                ("Điều 2: Trách nhiệm và đóng góp của các bên", 
                 "- Trách nhiệm đóng góp của Bên A: {{trach_nhiem_ben_a}}.\n- Trách nhiệm đóng góp của Bên B: {{trach_nhiem_ben_b}}."),
                ("Điều 3: Phân chia doanh thu và chi phí phát sinh", 
                 "- Cơ chế tài chính, phân bổ chi phí hoạt động và phân chia doanh thu cụ thể: {{co_che_tai_chinh}}."),
                ("Điều 4: Quyền sở hữu trí tuệ và tài sản thương hiệu", 
                 "- Mọi nhãn hiệu, logo và tài sản trí tuệ của bên nào vẫn thuộc quyền sở hữu riêng của bên đó.\n- Việc sử dụng thương hiệu chung cho chương trình liên kết phải được phê duyệt trước."),
                ("Điều 5: Thời hạn liên kết và điều khoản chấm dứt", 
                 "- Hợp đồng liên kết có thời hạn: {{thoi_han_lien_ket}}.\n- Chấm dứt trước hạn bằng văn bản thông báo trước tối thiểu {{thoi_gian_bao_truoc}} ngày.")
            ]
        },
        "Hop_dong_tin_dung": {
            "title": "HỢP ĐỒNG TÍN DỤNG",
            "clauses": [
                ("Điều 1: Hạn mức tín dụng và mục đích vay vốn", 
                 "- Bên A (Tổ chức tín dụng) cấp cho Bên B khoản vay hạn mức: {{han_muc_tin_dung}} VNĐ.\n- Mục đích sử dụng vốn vay của Bên B: {{muc_dich_vay}}."),
                ("Điều 2: Thời hạn cho vay và giải ngân", 
                 "- Thời hạn cho vay tối đa là: {{thoi_han_vay}} tháng.\n- Phương thức giải ngân tiền vay: chuyển khoản hoặc tiền mặt theo từng đề nghị rút vốn."),
                ("Điều 3: Lãi suất cho vay và các loại phí", 
                 "- Lãi suất cho vay trong hạn: {{lai_suat_trong_han}} %/năm.\n- Kỳ điều chỉnh lãi suất định kỳ: {{ky_dieu_chinh}} tháng/lần.\n- Phí thu xếp vốn và các chi phí quản lý tín dụng khác: {{phi_tin_dung}}."),
                ("Điều 4: Phương thức trả nợ gốc và lãi vay", 
                 "- Trả nợ lãi vay định kỳ hàng tháng vào ngày {{ngay_tra_lai}}.\n- Trả nợ gốc theo phân kỳ trả nợ cụ thể hoặc trả một lần vào ngày đáo hạn hợp đồng."),
                ("Điều 5: Các biện pháp bảo đảm nghĩa vụ nợ", 
                 "- Nghĩa vụ nợ của Bên B được bảo đảm bằng tài sản thế chấp / cầm cố / bảo lãnh theo Hợp đồng bảo đảm số: {{so_hop_dong_bao_dam}}.")
            ]
        },
        "Hop_dong_bao_lanh": {
            "title": "HỢP ĐỒNG BẢO LÃNH",
            "clauses": [
                ("Điều 1: Phạm vi bảo lãnh thực hiện nghĩa vụ", 
                 "- Bên A (Bên bảo lãnh) cam kết thực hiện nghĩa vụ thay cho {{ten_ben_duoc_bao_lanh}} nếu bên này không thực hiện đúng nghĩa vụ đối với Bên B."),
                ("Điều 2: Giá trị bảo lãnh tối đa", 
                 "- Số tiền bảo lãnh nghĩa vụ tối đa là: {{so_tien_bao_lanh}} VNĐ.\n- Nghĩa vụ được bảo lãnh bao gồm: {{nghia_vu_bao_lanh}} (Nghĩa vụ trả nợ, bàn giao...)."),
                ("Điều 3: Trách nhiệm thanh toán của Bên bảo lãnh", 
                 "- Trong vòng {{han_thanh_toan}} ngày kể từ khi nhận văn bản yêu cầu thực hiện nghĩa vụ bảo lãnh của Bên B, Bên A có trách nhiệm thanh toán số tiền bảo lãnh."),
                ("Điều 4: Quyền yêu cầu bồi hoàn và hoàn trả", 
                 "- Sau khi Bên A thực hiện nghĩa vụ bảo lãnh thay cho Bên được bảo lãnh, Bên A có quyền yêu cầu Bên được bảo lãnh hoàn trả toàn bộ số tiền cùng lãi phát sinh."),
                ("Điều 5: Cam kết chung và giải quyết tranh chấp", 
                 "- Cam kết bảo lãnh là vô điều kiện và không hủy ngang trong suốt thời hạn hiệu lực của bảo lãnh.")
            ]
        },
        "Hop_dong_the_chap": {
            "title": "HỢP ĐỒNG THẾ CHẤP TÀI SẢN",
            "clauses": [
                ("Điều 1: Tài sản thế chấp bảo đảm nghĩa vụ", 
                 "- Bên A đồng ý thế chấp tài sản thuộc quyền sở hữu hợp pháp của mình để bảo đảm nghĩa vụ: {{tai_san_the_chap}}.\n- Giấy tờ chứng nhận pháp lý kèm theo tài sản: {{giay_to_phap_ly}}."),
                ("Điều 2: Nghĩa vụ được bảo đảm bằng thế chấp", 
                 "- Bảo đảm cho nghĩa vụ trả nợ gốc, lãi và các chi phí phát sinh theo Hợp đồng vay vốn tín dụng số: {{so_hop_dong_vay}}."),
                ("Điều 3: Quyền và nghĩa vụ của Bên thế chấp (Bên A)", 
                 "- Được giữ và sử dụng tài sản thế chấp (hoặc bàn giao cho bên thứ ba quản lý).\n- Không được bán, trao đổi, tặng cho hoặc cho thuê tài sản thế chấp khi chưa có sự đồng ý của Bên B."),
                ("Điều 4: Quyền và nghĩa vụ của Bên nhận thế chấp (Bên B)", 
                 "- Giữ bản chính các giấy tờ pháp lý liên quan đến tài sản thế chấp.\n- Yêu cầu xử lý tài sản thế chấp để thu hồi nợ khi Bên A vi phạm nghĩa vụ thanh toán."),
                ("Điều 5: Xử lý tài sản thế chấp để thu hồi nợ", 
                 "- Phương thức xử lý tài sản thế chấp: bán đấu giá hoặc Bên B nhận chính tài sản để thay thế nghĩa vụ thanh toán theo thỏa thuận.")
            ]
        },
        "Hop_dong_cam_co": {
            "title": "HỢP ĐỒNG CẦM CỐ TÀI SẢN",
            "clauses": [
                ("Điều 1: Tài sản cầm cố và giao nhận tài sản", 
                 "- Bên A bàn giao tài sản cầm cố thuộc quyền sở hữu của mình cho Bên B nắm giữ: {{tai_san_cam_co}}.\n- Tình trạng tài sản thực tế và biên bản giao nhận kèm theo hợp đồng."),
                ("Điều 2: Nghĩa vụ được bảo đảm bằng cầm cố", 
                 "- Bảo đảm cho nghĩa vụ thực hiện trả nợ theo hợp đồng vay mượn tài sản số: {{so_hop_dong_vay}}."),
                ("Điều 3: Trách nhiệm bảo quản tài sản cầm cố của Bên B", 
                 "- Bên B có trách nhiệm giữ gìn, bảo quản tài sản cầm cố nguyên vẹn.\n- Không được sử dụng hoặc cho bên thứ ba thuê tài sản cầm cố nếu Bên A không đồng ý."),
                ("Điều 4: Quyền và nghĩa vụ của Bên cầm cố (Bên A)", 
                 "- Tự chịu mọi chi phí hợp lý để bảo quản và đăng ký giao dịch bảo đảm tài sản cầm cố.\n- Được nhận lại tài sản cầm cố sau khi hoàn thành đầy đủ nghĩa vụ nợ."),
                ("Điều 5: Xử lý tài sản cầm cố khi vi phạm nghĩa vụ", 
                 "- Nếu Bên A không trả nợ đúng hạn, Bên B có quyền bán hoặc chuyển nhượng tài sản cầm cố để cấn trừ nợ gốc và lãi.")
            ]
        },
        "Hop_dong_chuyen_giao_cong_nghe": {
            "title": "HỢP ĐỒNG CHUYỂN GIAO CÔNG NGHỆ",
            "clauses": [
                ("Điều 1: Đối tượng công nghệ được chuyển giao", 
                 "- Bên A đồng ý chuyển giao quyền sở hữu / quyền sử dụng công nghệ sản xuất: {{ten_cong_nghe}} cho Bên B.\n- Danh mục tài liệu kỹ thuật, sơ đồ quy trình kèm theo hợp đồng."),
                ("Điều 2: Phạm vi và hình thức chuyển giao công nghệ", 
                 "- Hình thức chuyển giao: {{hinh_thuc_chuyen_giao}} (chuyển giao độc quyền hoặc không độc quyền).\n- Lãnh thổ và thời hạn sử dụng công nghệ được chuyển giao: {{pham_vi_su_dung}}."),
                ("Điều 3: Giá chuyển giao và phương thức thanh toán", 
                 "- Phí chuyển giao công nghệ trọn gói là: {{gia_tri_chuyen_giao}} VNĐ.\n- Phương thức thanh toán tiền chuyển giao công nghệ: {{phuong_thuc_thanh_toan}}."),
                ("Điều 4: Đào tạo kỹ thuật và hỗ trợ vận hành", 
                 "- Bên A có trách nhiệm đào tạo nhân sự cho Bên B vận hành công nghệ trơn tru trong thời hạn {{thoi_gian_ho_tro}} ngày.\n- Cung cấp dịch vụ bảo hành và nâng cấp công nghệ định kỳ."),
                ("Điều 5: Quyền sở hữu trí tuệ và cam kết bảo mật", 
                 "- Bên B cam kết không tiết lộ bí mật công nghệ cho bất kỳ bên thứ ba nào.\n- Mọi cải tiến công nghệ do Bên B tự thực hiện trong quá trình sử dụng thuộc sở hữu của Bên B.")
            ]
        },
        "Hop_dong_nhuong_quyen_thuong_mai": {
            "title": "HỢP ĐỒNG NHƯỢNG QUYỀN THƯƠNG MẠI",
            "clauses": [
                ("Điều 1: Quyền thương mại được nhượng quyền", 
                 "- Bên A đồng ý cấp quyền thương mại cho Bên B sử dụng hệ thống nhãn hiệu, logo và mô hình cửa hàng: {{ten_thuong_hieu}} để kinh doanh."),
                ("Điều 2: Phạm vi lãnh thổ nhượng quyền và thời hạn hiệu lực", 
                 "- Khu vực địa lý độc quyền kinh doanh nhượng quyền của Bên B: {{dia_diem_kinh_doanh}}.\n- Thời hạn nhượng quyền thương mại là: {{thoi_han_nhuong_quyen}} năm."),
                ("Điều 3: Phí nhượng quyền và phí duy trì hoạt động", 
                 "- Phí nhượng quyền ban đầu (Franchise fee) là: {{phi_nhuong_quyen_ban_dau}} VNĐ.\n- Phí duy trì thương hiệu định kỳ (Royalty fee) hàng tháng: {{ty_le_phi_duy_tri}} % trên tổng doanh thu."),
                ("Điều 4: Tiêu chuẩn chất lượng và hỗ trợ vận hành", 
                 "- Bên A cung cấp bộ cẩm nang vận hành cửa hàng, thiết kế mẫu mã và hỗ trợ quảng bá thương hiệu chung.\n- Bên B phải tuân thủ nghiêm ngặt các quy định về chất lượng sản phẩm và dịch vụ của Bên A."),
                ("Điều 5: Quyền sở hữu trí tuệ và cam kết bảo mật thông tin", 
                 "- Bên B thừa nhận Bên A là chủ sở hữu duy nhất đối với thương hiệu và bí quyết kinh doanh nhượng quyền.\n- Cam kết bảo mật mọi thông tin tài chính và quy trình vận hành.")
            ]
        },
        "Hop_dong_so_huu_tri_tue": {
            "title": "HỢP ĐỒNG SỞ HỮU TRÍ TUỆ",
            "clauses": [
                ("Điều 1: Đối tượng quyền sở hữu trí tuệ chuyển nhượng", 
                 "- Bên A chuyển nhượng quyền sở hữu đối với nhãn hiệu / kiểu dáng công nghiệp / bản quyền tác giả: {{ten_doi_tuong_shtt}} cho Bên B.\n- Số Giấy chứng nhận đăng ký do Cục Sở hữu trí tuệ cấp: {{so_van_bang_bao_ho}}."),
                ("Điều 2: Giá trị chuyển nhượng và phương thức thanh toán", 
                 "- Giá chuyển nhượng đối tượng sở hữu trí tuệ là: {{gia_tri_chuyen_nhuong}} VNĐ.\n- Phương thức thanh toán tiền chuyển nhượng: {{phuong_thuc_thanh_toan}}."),
                ("Điều 3: Thủ tục chuyển quyền sở hữu pháp lý", 
                 "- Bên A cam kết hỗ trợ Bên B hoàn tất các thủ tục đăng ký chuyển nhượng văn bằng bảo hộ tại cơ quan quản lý nhà nước.\n- Chi phí, lệ phí thực hiện đăng ký do Bên {{ben_chiu_le_phi}} chi trả."),
                ("Điều 4: Quyền và nghĩa vụ của Bên chuyển nhượng", 
                 "- Cam kết đối tượng sở hữu trí tuệ thuộc quyền sở hữu độc quyền và hợp pháp của Bên A.\n- Ngừng sử dụng đối tượng chuyển nhượng kể từ ngày hợp đồng có hiệu lực đăng ký."),
                ("Điều 5: Quyền và nghĩa vụ của Bên nhận chuyển nhượng", 
                 "- Được toàn quyền sử dụng, khai thác thương mại và thực hiện các quyền nhân thân/tài sản pháp luật quy định đối với đối tượng nhận chuyển nhượng.")
            ]
        }
    }

    # Bắt đầu ghi log
    with open(log_path, "w", encoding="utf-8") as log_file:
        log_file.write("Bắt đầu tái tạo 17 tệp tin template hợp đồng chuẩn...\n")
        
        for filename, data in contracts_data.items():
            try:
                doc = Document()
                
                # Định dạng font chuẩn Times New Roman
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
                
                # 1. Quốc hiệu
                p_national = doc.add_paragraph()
                p_national.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_nat1 = p_national.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
                run_nat1.bold = True
                run_nat1.font.size = Pt(13)
                run_nat2 = p_national.add_run("Độc lập - Tự do - Hạnh phúc\n")
                run_nat2.bold = True
                run_nat2.font.size = Pt(14)
                run_nat3 = p_national.add_run("---------------")
                
                # Khoảng trống
                doc.add_paragraph()
                
                # 2. Tiêu đề Hợp đồng
                p_title = doc.add_paragraph()
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_title = p_title.add_run(data["title"])
                run_title.bold = True
                run_title.font.size = Pt(16)
                
                # Số hiệu hợp đồng
                p_no = doc.add_paragraph()
                p_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_no.add_run("Số: ....../HĐ")
                
                doc.add_paragraph()
                
                # 3. Căn cứ và thông tin các bên
                doc.add_paragraph("- Căn cứ Bộ luật Dân sự nước Cộng hòa xã hội chủ nghĩa Việt Nam hiện hành;\n- Căn cứ nhu cầu và khả năng thực tế của hai bên;")
                
                p_date = doc.add_paragraph()
                p_date.add_run("Hôm nay, ngày ...... tháng ...... năm ......, chúng tôi gồm:")
                
                p_ben_a = doc.add_paragraph()
                r_ben_a = p_ben_a.add_run("BÊN A: {{ten_ben_a}}\n")
                r_ben_a.bold = True
                p_ben_a.add_run("- Địa chỉ: {{dia_chi_ben_a}}\n- Mã số thuế / Số CCCD: {{mst_ben_a}}\n- Đại diện: {{dai_dien_ben_a}}        Chức vụ: {{chuc_vu_ben_a}}\n- Điện thoại: {{dien_thoai_ben_a}}")
                
                p_ben_b = doc.add_paragraph()
                r_ben_b = p_ben_b.add_run("BÊN B: {{ten_ben_b}}\n")
                r_ben_b.bold = True
                p_ben_b.add_run("- Địa chỉ: {{dia_chi_ben_b}}\n- Mã số thuế / Số CCCD: {{mst_ben_b}}\n- Đại diện: {{dai_dien_ben_b}}        Chức vụ: {{chuc_vu_ben_b}}\n- Điện thoại: {{dien_thoai_ben_b}}")
                
                p_intro = doc.add_paragraph()
                p_intro.add_run("Hai bên thống nhất thỏa thuận ký kết hợp đồng với các điều khoản cụ thể dưới đây:")
                
                # 4. Các điều khoản
                for clause_title, clause_content in data["clauses"]:
                    p_clause_t = doc.add_paragraph()
                    r_clause_t = p_clause_t.add_run(clause_title)
                    r_clause_t.bold = True
                    
                    p_clause_c = doc.add_paragraph()
                    p_clause_c.add_run(clause_content)
                
                # 5. Phần ký tên cuối trang
                doc.add_paragraph()
                doc.add_paragraph("Hợp đồng được lập thành 02 (hai) bản có giá trị pháp lý như nhau, mỗi bên giữ 01 bản để thực hiện.")
                
                p_sign = doc.add_paragraph()
                p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_sign = p_sign.add_run("ĐẠI DIỆN BÊN A                                            ĐẠI DIỆN BÊN B\n(Ký, ghi rõ họ tên)                                      (Ký, ghi rõ họ tên)")
                run_sign.bold = True
                
                output_path = templates_dir / f"{filename}.docx"
                doc.save(str(output_path))
                log_file.write(f"-> Tạo thành công: {filename}.docx (Dung lượng: {output_path.stat().st_size} bytes)\n")
            except Exception as e:
                log_file.write(f"-> Lỗi khi tạo {filename}.docx: {e}\n")
                
        log_file.write("Hoàn tất tái tạo toàn bộ 17 tệp tin template hợp đồng!\n")

if __name__ == "__main__":
    build_templates()
