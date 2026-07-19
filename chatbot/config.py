import os
import pathlib
import logging
from dotenv import load_dotenv

# Phân tích cấu trúc tất cả các file mẫu hợp đồng hiện tại
try:
    from docx import Document
    log_path = pathlib.Path(r"d:\Project Save\chatbot law\scratch\all_templates_summary_new.txt")
    templates_dir = pathlib.Path(r"d:\Project Save\chatbot law\chatbot\templates")
    with open(log_path, "w", encoding="utf-8") as f:
        f.write("=== TỔNG HỢP NỘI DUNG HIỆN TẠI CỦA CÁC MẪU HỢP ĐỒNG ===\n\n")
        for filepath in templates_dir.glob("*.docx"):
            f.write(f"Tên tệp: {filepath.name}\n")
            try:
                doc = Document(str(filepath))
                f.write(f"  Tổng số đoạn văn: {len(doc.paragraphs)}\n")
                non_empty = [p.text.strip() for p in doc.paragraphs if p.text.strip()][:5]
                for idx, text in enumerate(non_empty):
                    f.write(f"  Đoạn #{idx}: [{text[:150]}]\n")
            except Exception as e:
                f.write(f"  Lỗi đọc file: {e}\n")
            f.write("-" * 50 + "\n\n")
except Exception as ex:
    pass

logger = logging.getLogger("chatbot.config")

BASE_DIR = pathlib.Path(__file__).parent
DATA_DIR = BASE_DIR / "data_laws"
INDEX_DIR = BASE_DIR / "index_laws"
CONTRACT_DIR = BASE_DIR / "contracts"
TEMP_UPLOAD_DIR = BASE_DIR.parent / "temp_uploads"
DB_PATH = str(BASE_DIR.parent / "feedback.db")

for p in [DATA_DIR, INDEX_DIR, CONTRACT_DIR, TEMP_UPLOAD_DIR]:
    p.mkdir(exist_ok=True)

load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    logger.warning("GEMINI_API_KEY is not set! LLM features will not work.")
GCS_BUCKET_NAME = os.getenv("GCS_BUCKET_NAME") 
GCS_LAWS_PREFIX = os.getenv("GCS_LAWS_PREFIX", "law/")

EMBED_MODEL_NAME = "keepitreal/vietnamese-sbert"
QDRANT_HOST = os.getenv("QDRANT_HOST", "localhost")
QDRANT_PORT = int(os.getenv("QDRANT_PORT", "6333"))
QDRANT_COLLECTION = "vietnamese_laws"

LLM_PROVIDER = os.getenv("LLM_PROVIDER", "gemini")  # "gemini", "ollama", hoặc "openrouter"
OLLAMA_API_BASE = os.getenv("OLLAMA_API_BASE", "http://localhost:11434")
LOCAL_MODEL_NAME = os.getenv("LOCAL_MODEL_NAME", "qwen2.5:7b")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "qwen/qwen-2.5-7b-instruct:free")

